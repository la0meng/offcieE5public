import random
from docx import Document
from docx.shared import Pt

def generate_addition_problem(min_value, max_value):
    num1 = random.randint(min_value, max_value)
    num2 = random.randint(min_value, max_value)
    problem = f"{num1} + {num2}"
    answer = num1 + num2
    return problem, answer

# 添加其他口算类型的生成函数，根据需要扩展

def generate_problems_by_user_choice():
    # 用户选择口算类型
    print("请选择口算类型：")
    print("1. 加法")
    print("2. 减法")
    # 添加其他口算类型的选项

    choice = input("请输入选项（例如，输入1选择加法）: ")

    # 用户输入生成题目的数量
    num_problems = int(input("请输入要生成的题目数量: "))

    # 设置题目范围
    min_value = 1
    max_value = 100

    # 创建 Word 文档
    doc = Document()
    doc.add_heading('口算题目', level=1)

    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)

    # 根据用户选择生成题目
    for _ in range(num_problems):
        if choice == '1':
            problem, answer = generate_addition_problem(min_value, max_value)
        # 添加其他口算类型的判断并调用相应的生成函数

        # 添加问题段落
        problem_paragraph = doc.add_paragraph()
        problem_run = problem_paragraph.add_run(f"问题: {problem}")
        problem_run.bold = True  # 设置为粗体

        # 添加答案段落
        answer_paragraph = doc.add_paragraph()
        answer_run = answer_paragraph.add_run(f"答案: {answer}\n")
        answer_run.italic = True  # 设置为斜体

    # 保存 Word 文档
    doc.save('口算题目.docx')

if __name__ == "__main__":
    generate_problems_by_user_choice()
